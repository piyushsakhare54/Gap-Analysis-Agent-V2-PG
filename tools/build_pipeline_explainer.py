from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs")
DOCX_PATH = OUT_DIR / "requirements_gap_pipeline_interview_guide.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "1F2937"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_executive_summary(doc)
    add_architecture_diagram(doc)
    add_stage_walkthrough(doc)
    add_config_modes(doc)
    add_interview_script(doc)
    add_strengths_and_limits(doc)
    add_commands(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Requirements Gap Analysis Agent")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End-to-End Pipeline Interview Guide")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(90, 90, 90)

    add_callout(
        doc,
        "One-line explanation",
        "This system reads business and engineering transcripts, extracts requirements and solutions, retrieves the most relevant implementation evidence with open-source embeddings, then uses a local open-source LLM such as Qwen3-32B to analyze and verify gaps.",
    )


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("1. What Problem This Solves", level=1)
    add_bullets(
        doc,
        [
            "Business teams describe what customers need in meeting transcripts.",
            "Engineering teams describe what was built, deferred, or left out.",
            "The agent compares both sides and produces a requirements gap report.",
            "Large mode scales this by chunking transcripts and comparing only focused requirement-solution pairs.",
        ],
    )


def add_architecture_diagram(doc: Document) -> None:
    doc.add_heading("2. High-Level Architecture", level=1)
    rows = [
        ("Business transcripts", "Parser", "Chunker", "Qwen requirement extraction"),
        ("Engineering transcripts", "Parser", "Chunker", "Qwen solution extraction"),
        ("Extracted candidates", "Deduplication", "BGE-M3 embeddings", "Top-k retrieval"),
        ("Requirement + candidate solutions", "Qwen focused gap analysis", "Qwen critic", "Report writers"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [2100, 2100, 2100, 2460])
    headers = table.rows[0].cells
    for idx, label in enumerate(["Input", "Step 1", "Step 2", "Step 3 / Output"]):
        set_cell_text(headers[idx], label, bold=True)
        shade_cell(headers[idx], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph("Read this as a pipeline: each row flows left to right, and the final row produces the report.")


def add_stage_walkthrough(doc: Document) -> None:
    doc.add_heading("3. Each Pipeline Stage in Simple Words", level=1)
    stages = [
        ("1", "Transcript loading", "Reads .txt, .md, and .jsonl files from business and engineering folders.", "Keeps source path, speaker, and line number so every extracted item can cite evidence."),
        ("2", "Chunking", "Splits long transcripts into overlapping windows, for example 40 lines with 5 lines overlap.", "Prevents context overflow and preserves nearby context across chunk boundaries."),
        ("3", "Requirement extraction", "Qwen reads each business chunk and returns structured requirements as JSON.", "Each requirement has an id, statement, quote, line range, constraints, notes, and confidence."),
        ("4", "Solution extraction", "Qwen reads each engineering chunk and returns structured solutions as JSON.", "It also captures deferred work, scope limits, and tech stack when present."),
        ("5", "Deduplication", "Near-duplicate requirements and solutions from overlapping chunks are merged.", "This avoids repeated report items caused by chunk overlap."),
        ("6", "Embedding and vector index", "BAAI/bge-m3 converts requirement and solution text into vectors.", "The vector index uses cosine similarity to find semantically related solutions."),
        ("7", "Top-k matching", "For each requirement, the system retrieves the top 5 candidate solutions.", "This reduces the amount of text sent to the LLM and improves focus."),
        ("8", "Focused gap analysis", "Qwen compares one requirement against only its retrieved candidate solutions.", "It decides whether the requirement is covered, partial, or unaddressed."),
        ("9", "Critic verification", "Qwen reviews candidate gaps and removes false positives.", "This is a second-pass quality control layer."),
        ("10", "Report writing", "The final GapReport is written as Markdown, JSON, or JSONL.", "Metadata shows mode, chunk count, embedding model, top-k, warnings, and provider audit."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [550, 1800, 3250, 3760])
    headers = table.rows[0].cells
    for idx, label in enumerate(["#", "Stage", "What happens", "Why it matters"]):
        set_cell_text(headers[idx], label, bold=True)
        shade_cell(headers[idx], LIGHT_BLUE)
    for stage in stages:
        cells = table.add_row().cells
        for idx, value in enumerate(stage):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_config_modes(doc: Document) -> None:
    doc.add_heading("4. Standard Mode vs Large Open-Source Mode", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [2100, 3600, 3660])
    headers = table.rows[0].cells
    for idx, label in enumerate(["Area", "Standard mode", "Large Qwen mode"]):
        set_cell_text(headers[idx], label, bold=True)
        shade_cell(headers[idx], LIGHT_GRAY)
    rows = [
        ("Command flag", "Default when --mode is omitted", "--mode large --config configs/qwen_local.yaml"),
        ("Best for", "Small/simple transcripts and local testing", "Larger transcript datasets and interview demo"),
        ("Extraction", "Heuristic fallback in default config", "Qwen3-32B through Ollama"),
        ("Embeddings", "Hashing fallback", "BAAI/bge-m3 through sentence-transformers"),
        ("Gap reasoning", "Deterministic rules", "Qwen focused comparison"),
        ("Failure behavior", "Always runnable", "Fails clearly if Ollama/Qwen is unavailable"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_interview_script(doc: Document) -> None:
    doc.add_heading("5. Interview Explanation Script", level=1)
    add_callout(
        doc,
        "30-second version",
        "I built a requirements gap analysis agent for long transcripts. It chunks business and engineering conversations, uses Qwen to extract structured requirements and solutions, deduplicates overlap, embeds the items with BGE-M3, retrieves top-k matching solutions for each requirement, and then asks Qwen to perform focused gap analysis and critic verification before writing a report.",
    )
    doc.add_heading("If the interviewer asks why chunking is needed", level=2)
    add_bullets(
        doc,
        [
            "LLMs have context limits, and long meeting transcripts can exceed them.",
            "Chunking keeps each model call small and reliable.",
            "Overlap prevents losing context at chunk boundaries.",
        ],
    )
    doc.add_heading("If the interviewer asks why embeddings are needed", level=2)
    add_bullets(
        doc,
        [
            "Without retrieval, every requirement would need to be compared with every solution.",
            "Embeddings let the system find semantically relevant candidate solutions quickly.",
            "The LLM only reasons over focused evidence, which is cheaper, faster, and easier to audit.",
        ],
    )
    doc.add_heading("If the interviewer asks what is open source", level=2)
    add_bullets(
        doc,
        [
            "Qwen3-32B is the local reasoning model through Ollama.",
            "BAAI/bge-m3 is the embedding model through sentence-transformers.",
            "The vector index is a local in-memory cosine index, so FAISS is optional.",
        ],
    )


def add_strengths_and_limits(doc: Document) -> None:
    doc.add_heading("6. Strengths, Tradeoffs, and Future Work", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [3000, 3000, 3360])
    headers = table.rows[0].cells
    for idx, label in enumerate(["Strength", "Tradeoff", "Future improvement"]):
        set_cell_text(headers[idx], label, bold=True)
        shade_cell(headers[idx], LIGHT_BLUE)
    rows = [
        ("Scales beyond one LLM context window", "More chunks means more model calls", "Batching and async processing"),
        ("Focused evidence improves explainability", "Retrieval can miss weakly worded matches", "Hybrid keyword + vector retrieval"),
        ("Structured schemas make outputs testable", "LLM JSON can still need validation", "Retry and repair loop for malformed JSON"),
        ("Local open-source stack protects privacy", "Qwen 32B needs strong hardware", "Model size profiles such as 8B, 14B, 32B"),
        ("Modular design is easy to extend", "No graph reasoning yet", "GNN or knowledge graph as future work"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_commands(doc: Document) -> None:
    doc.add_heading("7. Commands to Run", level=1)
    commands = [
        ("Install dependencies", "pip install -r requirements.txt"),
        ("Pull Qwen model", "ollama pull qwen3:32b"),
        ("Run large Qwen mode", "python main.py --mode large --config configs/qwen_local.yaml --business transcripts/business --engineering transcripts/engineering --output reports/large_qwen_report.md --format markdown"),
        ("Run tests", "python -m pytest -q"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2300, 7060])
    headers = table.rows[0].cells
    set_cell_text(headers[0], "Action", bold=True)
    set_cell_text(headers[1], "Terminal command", bold=True)
    shade_cell(headers[0], LIGHT_GRAY)
    shade_cell(headers[1], LIGHT_GRAY)
    for label, command in commands:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], command, font_name="Consolas", font_size=8.5)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Requirements Gap Analysis Agent - Interview Guide")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [9360])
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run(body)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run("- ")
        p.add_run(item)


def set_cell_text(cell, text: str, bold: bool = False, font_name: str = "Calibri", font_size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(INK)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_widths(table, widths: list[int]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    main()

